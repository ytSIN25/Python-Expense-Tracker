# Default packages
import random as rnd #Generate ID for expense
import datetime as dt #Save expense date
import re #Validate user input format
import math #For certain calculation purpose
import os #Check is "ExpenseTrackerRecord.csv" exist and to force directory
import subprocess #Install package purpose
import sys #Check for module
import calendar #For ploting labels
from time import sleep #To stop program at some point for user to read prompted text

# Force running directory to be this program
scriptDirectory = os.path.dirname(os.path.abspath(__file__))
os.chdir(scriptDirectory)

#REMEMBER TO DELETE AFTER FINISH
import pandas as pd
import matplotlib.pyplot as plt
import docx

#Import when exist, install when not ✅
def importInstall(packageName, importName=None):
    importName = importName or packageName
    try:
        return __import__(importName)
    except ImportError:
        print(f"{packageName} not found. Installing...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", packageName])
        return __import__(importName)

# Non-default packages ✅
pd = importInstall("pandas")
plt = importInstall("matplotlib")
import matplotlib.pyplot as plt
docx = importInstall("python-docx", "docx")
openpyxl = importInstall("openpyxl")

# Sets up dataframe template ✅
def createNewDF():
    return pd.DataFrame({
        "ID": [],
        "name": [],
        "amount": [],
        "category": [],
        "date": [],
        "notes": []
    })

# Deal with CSV presaved data ✅
def initialize():
    global df
    if not os.path.exists("ExpenseTrackerRecord.csv"):
        print("ExpenseTrackerRecord.csv does not exist.")
        print('Creating "ExpenseTrackerRecord.csv" at current directory...')
        sleep(2)
        with open("ExpenseTrackerRecord.csv", "x") as file:
            pass
    try:
        df = pd.read_csv("ExpenseTrackerRecord.csv")

        # if frame is setted up but data is empty
        if df.empty:
            print("The original data is empty, creating new dataframe...")
            sleep(1.5)
            df = createNewDF()
            updateCSV()

        # framework with data
        else: 
            print("There is something in the data. Do you want to overwrite(Y/n)")
            overwrite = input("> ")
            while overwrite not in ["Y","n"]:
                overwrite = input("Please enter Y or n:  ")
            if overwrite == "Y":
                df = createNewDF()
                updateCSV()
                print("CSV overwritten with blank data.")
                sleep(1)
            else: 
                df["date"] = pd.to_datetime(df["date"]).dt.normalize()
                print("The original data will be used.")
                sleep(1)
                
    # totaly empty file
    except pd.errors.EmptyDataError:
        print("The CSV is totaly blank. Creating new template...")
        sleep(1)
        df = createNewDF()
        updateCSV()

#Display the starting menu ✅
def displayMenu():
    """Displays the starting menu with title"""
    print()
    print("*"*120)
    #graffi text "Expense tracker"
    print(r"""

──────────────────────────────────────────────────────────────────────────────────────────────────────────────────────
─██████████████─████████──████████─██████████████─██████████████─██████──────────██████─██████████████─██████████████─
─██░░░░░░░░░░██─██░░░░██──██░░░░██─██░░░░░░░░░░██─██░░░░░░░░░░██─██░░██████████──██░░██─██░░░░░░░░░░██─██░░░░░░░░░░██─
─██░░██████████─████░░██──██░░████─██░░██████░░██─██░░██████████─██░░░░░░░░░░██──██░░██─██░░██████████─██░░██████████─
─██░░██───────────██░░░░██░░░░██───██░░██──██░░██─██░░██─────────██░░██████░░██──██░░██─██░░██─────────██░░██─────────
─██░░██████████───████░░░░░░████───██░░██████░░██─██░░██████████─██░░██──██░░██──██░░██─██░░██████████─██░░██████████─
─██░░░░░░░░░░██─────██░░░░░░██─────██░░░░░░░░░░██─██░░░░░░░░░░██─██░░██──██░░██──██░░██─██░░░░░░░░░░██─██░░░░░░░░░░██─
─██░░██████████───████░░░░░░████───██░░██████████─██░░██████████─██░░██──██░░██──██░░██─██████████░░██─██░░██████████─
─██░░██───────────██░░░░██░░░░██───██░░██─────────██░░██─────────██░░██──██░░██████░░██─────────██░░██─██░░██─────────
─██░░██████████─████░░██──██░░████─██░░██─────────██░░██████████─██░░██──██░░░░░░░░░░██─██████████░░██─██░░██████████─
─██░░░░░░░░░░██─██░░░░██──██░░░░██─██░░██─────────██░░░░░░░░░░██─██░░██──██████████░░██─██░░░░░░░░░░██─██░░░░░░░░░░██─
─██████████████─████████──████████─██████─────────██████████████─██████──────────██████─██████████████─██████████████─
──────────────────────────────────────────────────────────────────────────────────────────────────────────────────────
────────────────────────────────────────────────────────────────────────────────────────────────────────────────────
─██████████████─████████████████───██████████████─██████████████─██████──████████─██████████████─████████████████───
─██░░░░░░░░░░██─██░░░░░░░░░░░░██───██░░░░░░░░░░██─██░░░░░░░░░░██─██░░██──██░░░░██─██░░░░░░░░░░██─██░░░░░░░░░░░░██───
─██████░░██████─██░░████████░░██───██░░██████░░██─██░░██████████─██░░██──██░░████─██░░██████████─██░░████████░░██───
─────██░░██─────██░░██────██░░██───██░░██──██░░██─██░░██─────────██░░██──██░░██───██░░██─────────██░░██────██░░██───
─────██░░██─────██░░████████░░██───██░░██████░░██─██░░██─────────██░░██████░░██───██░░██████████─██░░████████░░██───
─────██░░██─────██░░░░░░░░░░░░██───██░░░░░░░░░░██─██░░██─────────██░░░░░░░░░░██───██░░░░░░░░░░██─██░░░░░░░░░░░░██───
─────██░░██─────██░░██████░░████───██░░██████░░██─██░░██─────────██░░██████░░██───██░░██████████─██░░██████░░████───
─────██░░██─────██░░██──██░░██─────██░░██──██░░██─██░░██─────────██░░██──██░░██───██░░██─────────██░░██──██░░██─────
─────██░░██─────██░░██──██░░██████─██░░██──██░░██─██░░██████████─██░░██──██░░████─██░░██████████─██░░██──██░░██████─
─────██░░██─────██░░██──██░░░░░░██─██░░██──██░░██─██░░░░░░░░░░██─██░░██──██░░░░██─██░░░░░░░░░░██─██░░██──██░░░░░░██─
─────██████─────██████──██████████─██████──██████─██████████████─██████──████████─██████████████─██████──██████████─
────────────────────────────────────────────────────────────────────────────────────────────────────────────────────

""")
    author = "by Yee Tng-37289632"
    print(f"{author:>120}")
    print("*"*120)
    print("Welcome to the Expense Tracker Program")
    print("What do you like to do today?")
    print("1. View expenses")
    print("2. Add new expense")
    print("3. Edit past expense")
    print("4. Delete expense")
    print("5. Search for expense")
    print("6. View summary")
    print("7. Load CSV data")
    print("8. Clear CSV data")
    print("9. Export data")
    print("h for Help")
    print("Exit (q to exit)")
    print()
    print()

# Display expense list ✅
def displayExpense():
    print("""
    Do you want to see:
    1. All past expenses
    2. Latest 5 expenses
    3. All expenses in month
    (q to leave)
          """)
    choice = input("> ")
    while choice not in ["1", "2", "3","q","Q"]:
        choice = input("The choice is invalid:  ")
    match choice:
        case "1":
            df.sort_values(by="date", inplace=True) 
            df.reset_index(drop=True, inplace=True)
            print("*"*120)
            print(df)
            print("*"*120)
        case "2":
            df.sort_values(by="date", inplace=True) 
            df.reset_index(drop=True, inplace=True)
            print("*"*120)
            print(df.tail(5))
            print("*"*120)
        case "3":
            # Make sure date column is date
            df["date"] = pd.to_datetime(df["date"]).dt.normalize()

            Year = input("What is the year you would like to look for? ")
            validYear = df["date"].apply(lambda d: d.year).unique().tolist()
            while not Year.isdigit() or int(Year) not in validYear:
                Year = input("The year is not valid, try again: ")
            Year = int(Year)

            Month = input("What is the month you would like to look for? ")
            validMonth = df["date"].apply(lambda d: d.month).unique().tolist()
            while not Month.isdigit() or int(Month) not in validMonth:
                Month = input("The month is not valid, try again: ")
            Month = int(Month)

            filtered = df[(df["date"].dt.year == Year) & (df["date"].dt.month == Month)]
            print("*"*120)
            print(filtered)
            print("*"*120)
            

        case "q" | "Q":
            return
    
# Add new expense ✅
def addExpense():
    # Name
    def name():
        x = input("Enter the name for the expense (q to quit):  ")
        if x == "q" or x == "Q":
            return None
        return x
        
    # Amount
    def amount():
        amount = input("Enter the amount for the expense (r to redo / q to quit):  ")
        amountPattern = r"^\d+(\.\d{2})$"
        if amount == "r" or amount == "R":
            return "redo"
        elif amount == "q" or amount == "Q":
            return None
        while not re.match(amountPattern, amount):
            amount = input("The input is not valid (Add .00 for whole value):  ")
        return float(amount)

    # Category
    def category():
        print("""Choose the category of the expense:
    1. Food and Drinks
    2. Commute
    3. Utilities
    4. Entertainment / Shopping
    5. Social Life
    (r to redo / q to quit)
        """)
        validCategory = ["1","2","3","4","5","r","R","q","Q"]
        category = input("> ")
        while category not in validCategory:
            category = input("The category is not valid. Try Again: ")
        match category:
            case "1":
                category = "Food and Drinks"
            case "2":
                category = "Commute"
            case "3":
                category = "Utilities"
            case "4":
                category = "Entertainment / Shopping"
            case "5":
                category = "Social Life"
            case "r" | "R":
                return "redo"
            case "q" | "Q":
                return None
        return category

    # Date
    def date():
        Year = input("Enter the year of the expense : ")
        while not re.match(r"^\d{4}$", Year):
            Year = input("Enter only 4 digit number: ")
        Year = int(Year)

        Month = input("Enter the month of the expense in number: ")
        while not (re.match(r"^\d{1,2}$", Month) and 1 <= int(Month) <= 12):
            Month = input("Please enter a valid month in number: ")
        Month = int(Month)

        Day = input("Enter the day of the expense in number: ")
        if Month in [1,3,5,7,8,10,12]:
            while not (re.match(r"^\d{1,2}$", Day) and 1 <= int(Day) <= 31):
                Day = input("Please enter a valid day in number: ")
        elif Month in [4,6,9,11]:
            while not (re.match(r"^\d{1,2}$", Day) and 1 <= int(Day) <= 30):
                Day = input("Please enter a valid day in number: ")
        elif Month == 2 and Year % 4 == 0:
            while not (re.match(r"^\d{1,2}$", Day) and 1 <= int(Day) <= 29):
                Day = input("Please enter a valid day in number: ")
        elif Month == 2 and Year % 4 > 0:
            while not (re.match(r"^\d{1,2}$", Day) and 1 <= int(Day) <= 28):
                Day = input("Please enter a valid day in number: ")
        Day = int(Day)

        return pd.to_datetime(f"{Year}-{Month}-{Day}").normalize()

    # Main Logic
    while True:
        Name = name()
        if Name == None:
            return
        
        Amount = amount()
        if Amount == "redo":
            continue
        elif Amount is None:
            return

        Category = category()
        if Category == "redo":
            continue
        elif Category is None:
            return

        Date = date()

        # Note
        notes = input("Enter any extra notes. (If no, press Enter):  ")

        # ID generation
        id = genID(Category, Date.month, Date.day, Amount)

        df.loc[len(df)] = {"ID": id,
                        "name": Name,
                        "amount": Amount, 
                        "category": Category,
                        "date": Date,
                        "notes": notes}
        updateCSV()
        return
    
# Generate expense ID ✅
def genID(category, month, day, amount):
    #Category + Date + Random Number
    base = rnd.randint(0,100)
    head = 0
    match category:
        case "Food and Drinks":
            head  = 1
        case "Commute":
            head  = 2
        case "Utilities":
            head  = 3
        case "Entertainment / Shopping":
            head  = 4
        case "Social Life":
            head  = 5
    return int(str(head) + str(month) + str(day) + str(math.floor(amount)) + str(base))

# Edit expense ✅
def editExpense():
    target = input("Please enter the ID of the expense you would like to edit (q to quit):  ")
    while target not in df["ID"].astype(str).tolist() and target != "q":
        target = input("The ID is invalid. Try again:  ")
    if target == "q":
        return
    targetIndex = df.index[df["ID"].astype(str) == target][0]
    print("*"*120)
    print(f"The expense you would like to edit is: {df.at[targetIndex, 'name']} on {df.at[targetIndex, 'date']} with the amount {df.at[targetIndex, 'amount']} for {df.at[targetIndex, 'category']}. (Notes: {df.at[targetIndex, 'notes']})")
    print("""
    What is the data you would like to edit on this expense
    1. Name
    2. Date
    3. Amount
    4. Category
    5. Notes
    (q to quit)
          """)
    
    editTarget = input("> ")
    while editTarget not in ["1","2","3","4","5","q","Q"]:
        editTarget = input("The data is invalid. Try Again:  ")

    def confirm():
        confirm = input("Are you sure? (Y / n)  ")
        while confirm not in ["Y","n"]:
            confirm = input("Please enter Y(yes) or n(no):  ")
        if confirm == "Y":
            return True
        else:
            return False
        
    match editTarget:
        # Name
        case "1":
            name = input("Enter the name for the expense (q to quit):  ")
            if name == "q" or name == "Q":
                return
            if confirm():
                df.at[targetIndex, 'name'] = name
            else: 
                return

        # Date
        case "2":
            Year = input("Enter the year of the expense : ")
            while not re.match(r"^\d{4}$", Year):
                Year = input("Enter only 4 digit number: ")
            Year = int(Year)

            Month = input("Enter the month of the expense in number: ")
            while not (re.match(r"^\d{1,2}$", Month) and 1 <= int(Month) <= 12):
                Month = input("Please enter a valid month in number: ")
            Month = int(Month)

            Day = input("Enter the day of the expense in number: ")
            if Month in [1,3,5,7,8,10,12]:
                while not (re.match(r"^\d{1,2}$", Day) and 1 <= int(Day) <= 31):
                    Day = input("Please enter a valid day in number: ")
            elif Month in [4,6,9,11]:
                while not (re.match(r"^\d{1,2}$", Day) and 1 <= int(Day) <= 30):
                    Day = input("Please enter a valid day in number: ")
            elif Month == 2 and Year % 4 == 0:
                while not (re.match(r"^\d{1,2}$", Day) and 1 <= int(Day) <= 29):
                    Day = input("Please enter a valid day in number: ")
            elif Month == 2 and Year % 4 > 0:
                while not (re.match(r"^\d{1,2}$", Day) and 1 <= int(Day) <= 28):
                    Day = input("Please enter a valid day in number: ")
            Day = int(Day)

            if confirm():
                df.at[targetIndex, 'date'] = pd.to_datetime(f"{Year}-{Month}-{Day}").normalize()
                df.at[targetIndex, 'ID'] = genID(df.at[targetIndex, 'category'], df.at[targetIndex, 'date'].month, df.at[targetIndex, 'date'].day, df.at[targetIndex, 'amount'])
            else:
                return

        # Amount
        case "3":
            amount = input("Enter the amount for the expense (q to quit):  ")
            amountPattern = r"^\d+(\.\d{2})$"
            if amount == "q" or amount == "Q":
                return 
            while not re.match(amountPattern, amount):
                amount = input("The input is not valid (Add .00 for whole value):  ")
            if confirm():
                df.at[targetIndex, 'amount'] = float(amount)
                df.at[targetIndex, 'ID'] = genID(df.at[targetIndex, 'category'], df.at[targetIndex, 'date'].month, df.at[targetIndex, 'date'].day, df.at[targetIndex, 'amount'])
            else: 
                return

        # Category
        case "4":
            print("""Choose the category of the expense:
    1. Food and Drinks
    2. Commute
    3. Utilities
    4. Entertainment / Shopping
    5. Social Life
    (q to quit)
        """)
            validCategory = ["1","2","3","4","5","q","Q"]
            category = input("> ")
            while category not in validCategory:
                category = input("The category is not valid. Try Again: ")
            match category:
                case "1":
                    category = "Food and Drinks"
                case "2":
                    category = "Commute"
                case "3":
                    category = "Utilities"
                case "4":
                    category = "Entertainment / Shopping"
                case "5":
                    category = "Social Life"
                case "q" | "Q":
                    return
            if confirm():
                df.at[targetIndex, 'category'] = category
                df.at[targetIndex, 'ID'] = genID(df.at[targetIndex, 'category'], df.at[targetIndex, 'date'].month, df.at[targetIndex, 'date'].day, df.at[targetIndex, 'amount'])
            else: 
                return

        # Notes
        case "5":
            notes = input("Enter any extra notes. (If no, press Enter):  ")
            if confirm():
                df.at[targetIndex, 'notes'] = notes
            else:
                return

        #Quit
        case "q" | "Q":
            return
        
    sleep(0.5)
    print(f"The expense ({df.at[targetIndex, 'ID']}) is now: {df.at[targetIndex, 'name']} on {df.at[targetIndex, 'date']} with the amount {df.at[targetIndex, 'amount']} for {df.at[targetIndex, 'category']}. (Notes: {df.at[targetIndex, 'notes']})")
    updateCSV()

# Delete expense ✅
def delExpense():
    target = input("Please enter the ID of the expense you would like to delete (q to quit):  ")
    while target not in df["ID"].astype(str).tolist() and target != "q":
        target = input("The ID is invalid. Try again:  ")
    if target == "q":
        return
    targetIndex = df.index[df["ID"].astype(str) == target][0]
    print(f"The expense you would like to delete is: {df.at[targetIndex, 'name']} on {df.at[targetIndex, 'date']} with the amount {df.at[targetIndex, 'amount']}")

    # Confirmation 
    confirm = input("Are you confirm to delete the expense? (Y / n):  ")
    while confirm not in ["Y", "n"]:
        confirm = input("Please type Y (yes) or n (no)")
    if confirm == "Y":
        df.drop(targetIndex, inplace=True)
        updateCSV()
        print("Expense deleted successfully!")

# Search for expense ✅
def searchExpense():
    print("""
What parameter would you like to search with?
1. Name
2. Amount Range
3. Category
4. Date Range
(q to quit)
""")
    userInput = input("Choose the function or quit:  ")
    while userInput not in ["1","2","3","4","q","Q"]:
            userInput = input("The function is not available, choose again:  ")
    if userInput == "q" or userInput == "q":
        return
    
    match userInput:
        # Keyword search name
        case '1':
            target = input("Enter the name of the expense that you would like to search:  ").strip().lower()
            search = df[df["name"].str.lower().str.contains(target)]
            while search.empty:
                target = input("The name is not available. Try Again:  ")
                search = df[df["name"].str.lower().str.contains(target)]
            print("*"*120)
            print("This is the search results:  ")
            print(search)
            print("*"*120)

        # Amount Range
        case '2':
            target = input("Enter the amount range that you would like to search (minimum,maximum):  ")
            # \s* space, \d+ multiple digit, (\.\d+)? optional decimal
            while not re.match(r"^\s*\d+(\.\d+)?\s*,\s*\d+(\.\d+)?\s*$", target):
                target = input("The format is incorrect. Try again (minimum,maximum):  ")
            min,max = target.split(",")
            min = float(min)
            max = float(max)
            if min > max:
                min, max = max, min
            
            if min == max:    
                search = df[df["amount"] == min]
            else: 
                search = df[(df["amount"] >= min) & (df["amount"] <= max)]
            if search.empty:
                print("*"*120)
                print("There is no data for the amount range.")
                print("*"*120)
            else:
                search = search.sort_values(by="amount")
                print("*"*120)
                print("This is the search results:  ")
                print(search) 
                print("*"*120)

        # Category
        case '3':
            print("""Choose the category of the expense:
    1. Food and Drinks
    2. Commute
    3. Utilities
    4. Entertainment / Shopping
    5. Social Life
    (q to quit)
        """)
            validCategory = ["1","2","3","4","5","q","Q"]
            category = input("> ")
            while category not in validCategory:
                category = input("The category is not valid. Try Again: ")
            if category == "q" or category == "Q":
                return
            match category:
                case "1":
                    category = "Food and Drinks"
                case "2":
                    category = "Commute"
                case "3": 
                    category = "Utilities"
                case "4":
                    category = "Entertainment / Shopping"
                case "5":
                    category = "Social Life"

            search = df[df["category"] == category]
            
            if search.empty:
                print("There is no expense in this category")
            else:
                search = search.sort_values(by="date")
                sleep(1)
                print("*"*120)
                print("This is the search results:  ")
                print(search) 
                print("*"*120)
        
        # Date range
        case '4':
            while True:
                startDate = input("Please enter the start date as YYYY-MM-DD (q to quit): ")
                if startDate.lower() == "q":
                    break
                
                if not re.match(r"^\d{4}-\d{2}-\d{2}$", startDate):
                    print("Invalid format. Use YYYY-MM-DD.")
                    continue
                
                try:
                    year, month, day = map(int, startDate.split("-"))
                    dt.date(year, month, day)
                    break
                except ValueError:
                    print("Invalid date. Try again in YYYY-MM-DD")
            year, month, day = map(int, startDate.split("-"))
            startDate = pd.to_datetime(f"{year}-{month}-{day}").normalize()
            

            while True:
                endDate = input("Please enter the end date as YYYY-MM-DD (q to quit): ")
                if endDate.lower() == "q":
                    break
                
                if not re.match(r"^\d{4}-\d{2}-\d{2}$", endDate):
                    print("Invalid format. Use YYYY-MM-DD.")
                    continue
                
                try:
                    year, month, day = map(int, endDate.split("-"))
                    dt.date(year, month, day)
                    break
                except ValueError:
                    print("Invalid date. Try again in YYYY-MM-DD")
            year, month, day = map(int, endDate.split("-"))
            endDate = pd.to_datetime(f"{year}-{month}-{day}").normalize()

            if startDate > endDate:
                startDate, endDate = endDate, startDate
            
            if startDate == endDate:    
                search = df[df["date"] == startDate]
            else: 
                search = df[(df["date"] >= startDate) & (df["date"] <= endDate)]
            if search.empty:
                print("*"*120)
                print("There is no data for the amount range.")
                print("*"*120)
            else:
                search = search.sort_values(by="category")
                print("*"*120)
                print("This is the search results:  ")
                print(search)
                print("*"*120)
    
# View summary
def viewSummary():
    while True:
        print("""
    Select the summary you would like to see:
    1. Daily
    2. Monthly
    3. Yearly
    4. Category
    (q to quit)
    """)
        selectSummary = input("> ")
        while selectSummary not in ["1","2","3","4","q","Q"]:
            selectSummary = input("The selection is invalid. Try Again: ")
        if selectSummary == "q" or selectSummary == "Q":
            break
        match selectSummary:
            # Day
            case "1":
                while True:
                    searchDate = input("Please enter the date as YYYY-MM-DD (q to quit): ")
                    if searchDate.lower() == "q":
                        break
                    
                    if not re.match(r"^\d{4}-\d{2}-\d{2}$", searchDate):
                        print("Invalid format. Use YYYY-MM-DD.")
                        continue
                    
                    try:
                        year, month, day = map(int, searchDate.split("-"))
                        dt.date(year, month, day)
                        break
                    except ValueError:
                        print("Invalid date. Try again in YYYY-MM-DD")
                year, month, day = map(int, searchDate.split("-"))
                searchDate = pd.to_datetime(f"{year}-{month}-{day}").normalize()

                result = df[df["date"] == searchDate]
                if result.empty:
                    print("There is no data on that day.")
                else:
                    total = sum(result["amount"].tolist())
                    print(f"The total expense on {searchDate} is RM{total}.")

            # Month
            case "2":
                Year = input("Enter the year of the expense : ")
                while not re.match(r"^\d{4}$", Year):
                    Year = input("Enter only 4 digit number: ")
                Year = int(Year)

                Month = input("Enter the month of the expense in number: ")
                while not (re.match(r"^\d{1,2}$", Month) and 1 <= int(Month) <= 12):
                    Month = input("Please enter a valid month in number: ")
                Month = int(Month)

                result = df[(df["date"].dt.year == Year) & (df["date"].dt.month == Month)]
                if result.empty:
                    print("There is no data in the month.")
                else:
                    total = sum(result["amount"].tolist())
                    if Month in [1,3,5,7,8,10,12]:
                        mean = total/31
                    elif Month in [4,6,9,11]:
                        mean = total/30
                    elif Month == 2 and Year % 4 == 0:
                        mean = total/29
                    elif Month == 2 and Year % 4 > 0:
                        mean = total/28
                    result = result.sort_values("amount")
                    highestDay = result.iloc[-1]
                    lowestDay = result.iloc[0]
                    print(f"The total expense on {Month}/{Year} is RM{total}.")
                    print(f"The average daily expense of this month is RM{mean}.")
                    print(f"The highest expense in the month is {highestDay['name']} on {highestDay['date']} with an amount of RM{highestDay['amount']}.")
                    print(f"The lowest expense in the month is {lowestDay['name']} on {lowestDay['date']} with an amount of RM{lowestDay['amount']}.")
                    
            # Year
            case "3":
                Year = input("Enter the year of the expense : ")
                while not re.match(r"^\d{4}$", Year):
                    Year = input("Enter only 4 digit number: ")
                Year = int(Year)
                result = df[df["date"].dt.year == Year]
                if result.empty:
                    print("There is no data for the year.")
                else: 
                    total = sum(result["amount"].tolist())
                    if Year % 4 == 0:
                        mean = total / 366
                    elif Year % 4 > 0:
                        mean = total / 365
                    monthlyExpense = result.groupby(result["date"].dt.month)["amount"].sum().to_dict()
                    maxMonth = max(monthlyExpense, key=monthlyExpense.get)
                    maxAmount = monthlyExpense[maxMonth]
                    minMonth = min(monthlyExpense, key=monthlyExpense.get)
                    minAmount = monthlyExpense[minMonth]
                    print("*"*120)
                    print(f"The total expense for year {Year} is RM{total}.")
                    print(f"The month with the highet expense is {calendar.month_name[maxMonth]} with a total expense of RM{maxAmount}.")
                    print(f"The month with the lowest expense is {calendar.month_name[minMonth]} with a total expense of RM{minAmount}.")

                    graph = input("Do you want a graph for monthly expenses? (y / n)  ")
                    while graph not in ["y","Y","n"]:
                        graph = input("Please enter y(yes) or n(no):  ")
                    if graph == "y" or graph == "Y":
                        fig, ax = plt.subplots(figsize=(7.5, 7.5))
                        bars = ax.bar(monthlyExpense.keys(), monthlyExpense.values(), color='yellow')
                        ax.bar_label(bars, padding=2)
                        ax.set_xlabel("Month", fontweight="bold")
                        ax.set_ylabel("Total Expense (RM)", fontweight="bold")
                        ax.set_xticks([num for num in range(1,13)])
                        ax.set_xticklabels([calendar.month_abbr[m] for m in range(1,13)])
                        plt.show()
                    print("*"*120)
                    
            # Category
            case "4":
                pass

# Export data ✅
def export():
    global df
    def create(name,format):
        global df
        if os.path.exists(f"{name}.{format}"):
            print(f"The file with the name '{name}' already exist. please delete and try again.")
            sleep(0.5)
            return
        print(f"Exporting to {name}.{format}...")

        if format == "xlsx":
            df.to_excel(f"{name}.{format}",index=False)

        elif format == "csv":
            df.to_csv(f"{name}.{format}",index=False)

        elif format == "html":
            df.to_html(f"{name}.{format}",index=False)

        elif format == "pdf":
            fig, ax = plt.subplots(figsize=(12, len(df)*0.5))
            ax.axis('off')
            table = ax.table(cellText = df.values, colLabels = df.columns, cellLoc = 'center', loc = 'center')
            plt.savefig(f"{name}.pdf")

        elif format == "docx":
            doc = docx.Document()
            doc.add_heading(name, 1)
            table = doc.add_table(rows=1, cols=len(df.columns))
            for i, column in enumerate(df.columns):
                table.rows[0].cells[i].text = column
            
            for row in df.values:
                cells = table.add_row().cells
                for i, value in enumerate(row):
                    cells[i].text = str(value)
            
            doc.save(f"{name}.docx")
            
        elif format == "txt":
            with open(f"{name}.txt", "w") as f:
                f.write(df.to_string(index=False))
        sleep(1)
        print("The data has been exported! ")

    # Main
    print("""
What file format do you want to export your data to?
1. Excel(xlsx)
2. CSV
3. HTML
4. PDF
5. Word(docx)
6. txt
(q to quit)
          """)
    select = input("> ")
    while select not in ["1","2","3","4","5","6","q","Q"]:
        select = input("Please select a valid option:  ")
    if select == "q" or select ==  "Q":
        return
    name = input("Enter the name you would like to have: ")
    name = name.replace(" ", "_")
    match select:
        case "1":
            create(name, "xlsx")
        case "2":
            create(name, "csv")
        case "3":
            create(name, "html")
        case "4":
            create(name, "pdf")
        case "5":
            create(name, "docx")
        case "6":
            create(name, "txt")

#help ✅
def help():
    while True:
        print()
        print("*"*120)
        print("Welcome to help ")
        print("Which function do you want to query?")
        print("1. View expenses")
        print("2. Add new expense")
        print("3. Edit past expense")
        print("4. Delete expense")
        print("5. Search for expense")
        print("6. View summary")
        print("7. Load CSV data")
        print("8. Clear CSV data")
        print("9. Export data")
        print("Exit (q to exit)")

        userInput = input("Choose the function or quit:  ")
        while userInput not in ["1","2","3","4","5","6","7","8","9","q","Q"]:
                userInput = input("The function is not available, choose again:  ")
        if userInput == "q" or userInput == "Q":
            break
        
        print("*"*120)
        match userInput:
            case '1':
                print("""
View Expense function is to view datas.
1. All Past Expense: Displays all the expenses in data
2. Latest 5 Expense: Displays the most latest inputted 5 data
3. All Expenses by Month: Displays all expense in the month chosen
""")
            case '2':
                print("""
Add Expense function is to create new expense record into the data.
Fill in Name, amount, category and date. 
An ID for the expense will be generated automatically.
ID is used to locate the expense in other part of this program.
""")
            case '3':
                print("""
Edit Expense function is to edit data recorded (as the name suggests).
ID for the expense is needed so remember to view and copy it.
You can edit the name, amount, date and category.
If the amount, date or category is changed, the ID will also be updated.
""")
            case '4':
                print("""
Delete Expense function is to delete data recorded (ofc).
Same as edit, an ID is needed.
Before deleting, a comfirmation question will be asked.
""")
            case '5':
                print("""
Search Expense function is to search for specific data.
1. Name: Type any keyword to find all expenses with name including the keyword.
2. Amount Range: Type a range of amount. If the min and max is the same, the specific expense will be shown.
3. Category: Displays all the expenses in the category
4. Date Range: A range of date are use to filter data. If start and end date is the same all expense on that day will be shown.
""")
            case '6':
                print("""
View Summary function is to view statistics on your data.
1. Analysis in Range: Displays analysis in a day, a month or a year.
(total, mean, distributions)
2. Analysis by Category: Displays which category is highest and lowest, optional to show a pie chart.
""")
            case '7':
                print("""
Loads the "ExpenseTrackerRecord.csv" manually.
Make sure such file is the same directory with this program.
(Loading data into the program is usually automated when you start the program. Use this function when you have manually edited the CSV.)
""")
            case '8':
                print("Clears everything in the 'ExpenseRecordTracker.csv' leaving a blank template. ")
            case '9':
                print("Exports the data to be Excel, Word, PDF, CSV, TXT or even HTML")
        print("*"*120)
        sleep(5)

#update csv to df ✅
def loadCSV():
    df = pd.read_csv("ExpenseTrackerRecord.csv")
    df["date"] = pd.to_datetime(df["date"], errors='coerce').dt.normalize()
    return df

#update df to csv ✅
def updateCSV():
    df.sort_values(by="date", inplace = True)
    df.reset_index(drop = True, inplace = True)
    df.to_csv("ExpenseTrackerRecord.csv",index=False)

def main():
    global df
    isRunning = True
    #Initialize in case "ExpenseTrackerRecord.csv" not exist
    initialize()

    #Main part
    while isRunning:
        displayMenu()
        userInput = input("Choose the function or quit:  ")
        while userInput not in ["1","2","3","4","5","6","7","8","9","h","H","q","Q"]:
                userInput = input("The function is not available, choose again:  ")

        match userInput:
            # Display Expense
            case '1':
                displayExpense()
                input("Press enter to go back to the menu")

            # Add new Expense
            case '2':
                addExpense()
                input("Press enter to go back to the menu")

            # Edit Expense
            case '3':
                editExpense()
                input("Press enter to go back to the menu")

            # Delete Expense
            case '4':
                delExpense()
                input("Press enter to go back to the menu")

            case '5':
                searchExpense()
                input("Press enter to go back to the menu")
            case '6':
                viewSummary()
                input("Press enter to go back to the menu")
            case '7':
                df = loadCSV()
                df["date"] = pd.to_datetime(df["date"]).dt.normalize()
                print("The CSV from last time has been imported!")
                sleep(2)
            case '8':
                confirm = input("Are you sure? (Y / n)  ")
                while confirm not in ["Y","n"]:
                    confirm = input("Please enter Y(yes) or n(no):  ")
                if confirm == "Y":
                    df = pd.DataFrame({"ID":[],
                                       "name": [],
                                       "amount": [],
                                       "category": [],
                                       "date": [],
                                       "notes": []})
                    updateCSV()
                input("Press enter to go back to the menu")
            case '9':
                export()
                input("Press enter to go back to the menu")
            # Help
            case 'h' | 'H':
                help()
                input("Press enter to go back to the menu")

            # Exit
            case 'q' | 'Q':
                print("Bye Bye!")
                isRunning = False
        
if __name__ == "__main__":
    main()